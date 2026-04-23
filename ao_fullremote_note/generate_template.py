from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── ヘルパー関数 ──────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run_jp(paragraph, text, bold=False, size=None, color=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = '游明朝'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'■ {title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = '游明朝'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    # 下線（青）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_fill_box(doc, hint_lines, empty_lines=3, bg='F5F5F5'):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_border(cell)
    cell.width = Cm(16)

    first = True
    for hint in hint_lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(hint)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        run.font.name = '游明朝'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    for _ in range(empty_lines):
        ep = cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(6)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_job_table(doc, job_label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'【{job_label}】')
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = '游明朝'
    r._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

    rows_def = [
        ('在籍期間',    '　　年　月 〜 　　年　月',  1),
        ('会社名・規模', '会社名：　　　　　（業種：　　　/ 従業員数：　　名）', 1),
        ('担当業務',    '・\n・\n・', 3),
        ('実績・取り組み', '（数字・成果を入れて書こう）\n・\n・\n・', 4),
        ('使用ツール',  '（例：Excel / Google Workspace / Zoom / Slack 等）', 1),
        ('役職・規模',  '（例：一般職 / チーム○名）', 1),
    ]

    table = doc.add_table(rows=len(rows_def), cols=2)
    set_col_width(table, 0, Cm(3.2))
    set_col_width(table, 1, Cm(13))

    for i, (label, hint, lines) in enumerate(rows_def):
        lc = table.cell(i, 0)
        vc = table.cell(i, 1)
        set_cell_bg(lc, 'D9E1F2')
        set_cell_border(lc)
        set_cell_border(vc)

        lc.paragraphs[0].clear()
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.name = '游明朝'
        lr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        vc.paragraphs[0].clear()
        vr = vc.paragraphs[0].add_run(hint)
        vr.font.size = Pt(9)
        vr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        vr.italic = True
        vr.font.name = '游明朝'
        vr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
        for _ in range(lines - 1):
            vc.add_paragraph()

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def set_col_width(table, col_idx, width):
    for row in table.rows:
        row.cells[col_idx].width = width


# ── ドキュメント生成 ─────────────────────────

doc = Document()

# ページ設定
sec = doc.sections[0]
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)

# デフォルトスタイル
normal = doc.styles['Normal']
normal.font.name = '游明朝'
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# ── タイトル ──
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(12)
tr = title_p.add_run('職　務　経　歴　書')
tr.bold = True
tr.font.size = Pt(18)
tr.font.name = '游明朝'
tr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# 作成日・氏名
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date_p.paragraph_format.space_after = Pt(2)
dr = date_p.add_run('　　　　年　　月　　日現在')
dr.font.size = Pt(10)
dr.font.name = '游明朝'
dr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

name_p = doc.add_paragraph()
name_p.paragraph_format.space_after = Pt(8)
nr = name_p.add_run('氏名：')
nr.bold = True
nr.font.size = Pt(10.5)
nr.font.name = '游明朝'
nr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
nr2 = name_p.add_run('　　　　　　　　　　　　　　　　　　　　')
nr2.font.size = Pt(10.5)
nr2.font.name = '游明朝'
nr2._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# ── ① 職務要約 ──
add_section_title(doc, '職務要約')
add_fill_box(doc, [
    '（記入例）○○職として○年の経験を持ちます。主に○○を担当してまいりました。',
    'フルリモートという働き方でより高いパフォーマンスを発揮できると考え、転職活動を行っています。',
], empty_lines=4)

# ── ② 活かせる経験・知識・技術 ──
add_section_title(doc, '活かせる経験・知識・技術')
add_fill_box(doc, [
    '（3〜6項目を箇条書きで。「フルリモートで活かせる能力」を意識して書く）',
    '・Excelを使った集計・データ管理（関数・ピボットテーブル等）',
    '・請求書・経費精算の処理業務（月○件対応）',
    '・Zoom / Teams を使ったオンライン会議の運営・議事録作成',
], empty_lines=5)

# ── ③ 自己PR ──
add_section_title(doc, '自己ＰＲ')
note_p = doc.add_paragraph()
note_p.paragraph_format.space_after = Pt(4)
nr = note_p.add_run('フルリモートで活かせる強みを2〜3つ、エピソード付きで書こう')
nr.font.size = Pt(9)
nr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
nr.font.name = '游明朝'
nr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

pr_items = [
    ('① 自律した業務遂行と進捗管理',
     '（例）前職では上司が常に近くにいる環境ではなかったため、自分でタスクを整理・優先順位をつけて業務を進める習慣が身についています。毎朝ToDoをリスト化し、夕方には完了状況をチームに共有する運用を自ら提案・実施しました。'),
    ('② 迅速な報告・連絡・相談の徹底',
     '（例）業務上のトラブルや疑問が生じた際には、自己判断で抱え込まず、速やかに上司・関係者へ共有することを心がけてきました。フルリモートでは「今何をしているか」が伝わりにくい分、こまめな発信を意識しています。'),
    ('③ 業務の仕組み化・マニュアル整備',
     '（例）前職では引き継ぎ資料が存在しなかったため、自ら業務フローを整理し手順書を一から作成しました。これにより引き継ぎ期間を○週間から○日間に短縮できました。'),
]

for sub_title, hint in pr_items:
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(8)
    sp.paragraph_format.space_after = Pt(2)
    sr = sp.add_run(sub_title)
    sr.bold = True
    sr.font.size = Pt(10)
    sr.font.name = '游明朝'
    sr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    add_fill_box(doc, [hint], empty_lines=4)

# ── ④ 取得資格・使用ツール ──
add_section_title(doc, '取得資格・使用ツール')
add_fill_box(doc, [
    '【資格】　（取得年月）　（資格名）　／　例：2020年3月　日商簿記2級',
    '【PCスキル・ツール】　Microsoft Office / 会計ソフト（弥生・freee等）/ Zoom / Slack 等',
], empty_lines=4)

# ── ⑤ 職務経歴 ──
add_section_title(doc, '職務経歴')
add_job_table(doc, '直近の勤務先')
add_job_table(doc, '前職（あれば）')
add_job_table(doc, '前々職（あれば）')

# ── ⑥ ブランク期間 ──
add_section_title(doc, 'ブランク期間がある方（該当する場合のみ）')
add_fill_box(doc, [
    '（例）20○○年○月〜20○○年○月：育児のため休職',
    '現在は保育園・学童の体制が整っており、フルタイム就業が可能な状態です。',
    '休職期間中は○○の勉強・資格取得など自己研鑽に取り組みました。',
], empty_lines=3)

# ── チェックリスト ──
add_section_title(doc, '提出前チェックリスト')
checks = [
    '職務要約は3〜5行に収まっているか',
    '「担当しました」で終わらず、成果・数字が書けているか',
    'セクションの順番は　自己PR → 資格 → 職務経歴　になっているか',
    'フルリモートに関連するツール名が書けているか',
    '自己PRで「自律性・報告力・仕組み化」のどれかが伝わるか',
    'ブランク期間がある場合、一言説明が入っているか',
    '誤字脱字がないか（2回以上読み返したか）',
]
for item in checks:
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    cp.paragraph_format.left_indent = Cm(0.5)
    cr = cp.add_run(f'☐　{item}')
    cr.font.size = Pt(10)
    cr.font.name = '游明朝'
    cr._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

# 以上
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
end_p.paragraph_format.space_before = Pt(20)
er = end_p.add_run('以上')
er.font.size = Pt(10)
er.font.name = '游明朝'
er._r.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

output = 'C:/Users/henac/prompt_vault/ao_fullremote_note/職務経歴書テンプレート.docx'
doc.save(output)
print(f'生成完了: {output}')
